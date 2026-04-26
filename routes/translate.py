from flask import Blueprint, request, jsonify, send_from_directory, send_file
import os, uuid
from docx import Document
import pdfplumber

from services.translation_service import do_translate
from services.voice_service import create_voice
from services.audio_service import transcribe_audio

translate_bp = Blueprint("translate", __name__)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

BASE_URL = "https://procedure-upchuck-praying.ngrok-free.dev"


# =========================
# 🎤 AUDIO TRANSLATION
# =========================
@translate_bp.route("/translate", methods=["POST"])
def translate_audio():
    try:
        if "audio" not in request.files:
            return jsonify({"error": "No audio uploaded"}), 400

        audio = request.files["audio"]
        direction = request.form.get("direction")

        path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.webm")
        audio.save(path)

        try:
            text = transcribe_audio(path, direction)
        except:
            text = ""

        if not text.strip():
            return jsonify({"error": "No speech"}), 200

        translated = do_translate(text, direction)

        filename = f"{uuid.uuid4().hex}.mp3"
        out = os.path.join(UPLOAD_FOLDER, filename)

        create_voice(translated, direction, "female", out)

        return jsonify({
            "original": text,
            "translated": translated,
            "audio": f"{BASE_URL}/uploads/{filename}"
        })

    except Exception as e:
        print("AUDIO ERROR:", e)
        return jsonify({"error": "Audio failed"}), 200


# =========================
# ✍️ TEXT TRANSLATION
# =========================
@translate_bp.route("/translate-text", methods=["POST"])
def translate_text():
    try:
        data = request.get_json()

        text = data.get("text", "").strip()
        direction = data.get("direction")

        if not text:
            return jsonify({"error": "Empty text"}), 400

        translated = do_translate(text, direction)

        filename = f"{uuid.uuid4().hex}.mp3"
        out = os.path.join(UPLOAD_FOLDER, filename)

        create_voice(translated, direction, "female", out)

        return jsonify({
            "original": text,
            "translated": translated,
            "audio": f"{BASE_URL}/uploads/{filename}"
        })

    except Exception as e:
        print("TEXT ERROR:", e)
        return jsonify({"error": "Text failed"}), 500


# =========================
# 📄 DOCUMENT TRANSLATION
# =========================
@translate_bp.route("/translate-doc", methods=["POST"])
def translate_doc():
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["file"]
        direction = request.form.get("direction")

        input_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}_{file.filename}")
        file.save(input_path)

        extracted_text = ""

        if file.filename.lower().endswith(".pdf"):
            with pdfplumber.open(input_path) as pdf:
                for page in pdf.pages:
                    extracted_text += (page.extract_text() or "") + "\n"

        elif file.filename.lower().endswith(".docx"):
            doc = Document(input_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"

        else:
            return jsonify({"error": "Unsupported format"}), 400

        if not extracted_text.strip():
            return jsonify({"error": "No readable text"}), 400

        new_doc = Document()

        for line in extracted_text.split("\n"):
            if line.strip():
                try:
                    translated = do_translate(line, direction)
                except:
                    translated = line
                new_doc.add_paragraph(translated)
            else:
                new_doc.add_paragraph("")

        output = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.docx")
        new_doc.save(output)

        return send_file(output, as_attachment=True, download_name="translated.docx")

    except Exception as e:
        print("DOC ERROR:", e)
        return jsonify({"error": "Document failed"}), 500


# =========================
# 🔥 LIVE TRANSLATION (FIXED)
# =========================
@translate_bp.route("/live-translate", methods=["POST"])
def live_translate():
    try:
        if "audio" not in request.files:
            return jsonify({"error": "No audio"}), 200

        audio = request.files["audio"]
        direction = request.form.get("direction", "en-fr")

        path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.webm")
        audio.save(path)

        # 🎤 SAFE TRANSCRIBE
        try:
            text = transcribe_audio(path, direction)
        except Exception as e:
            print("TRANSCRIBE ERROR:", e)
            text = ""

        # ❌ skip empty audio
        if not text.strip():
            return jsonify({"error": "No speech"}), 200

        # 🌍 TRANSLATE
        try:
            translated = do_translate(text, direction)
        except Exception as e:
            print("TRANSLATE ERROR:", e)
            return jsonify({"error": "Translate failed"}), 200

        # 🔊 VOICE
        filename = f"{uuid.uuid4().hex}.mp3"
        out = os.path.join(UPLOAD_FOLDER, filename)

        try:
            create_voice(translated, direction, "female", out)
        except Exception as e:
            print("VOICE ERROR:", e)
            return jsonify({"error": "Voice failed"}), 200

        return jsonify({
            "text": text,
            "translated": translated,
            "audio": f"{BASE_URL}/uploads/{filename}"
        })

    except Exception as e:
        print("LIVE ERROR:", e)
        return jsonify({"error": "Live failed"}), 200


# =========================
# 📥 SERVE FILES
# =========================
@translate_bp.route("/uploads/<filename>")
def serve_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)